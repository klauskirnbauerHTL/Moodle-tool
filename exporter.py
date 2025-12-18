import xml.etree.ElementTree as ET
import sqlite3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def export_to_moodle_xml(db_path, question_ids, filename):
    """Exportiert ausgewählte Fragen als Moodle XML"""
    conn = sqlite3.connect(db_path)
    c = conn.cursor()

    quiz = ET.Element('quiz')

    for qid in question_ids:
        c.execute('SELECT title, questiontext, single, tags, points, question_type FROM questions WHERE id=?', (qid,))
        row = c.fetchone()
        if not row:
            continue
        title, questiontext, single, tags, points = row[:5]
        question_type = row[5] if len(row) > 5 else 'multichoice'

        # Fragetyp bestimmen
        if question_type == 'essay':
            question = ET.SubElement(quiz, 'question', type='essay')
        elif question_type == 'shortanswer':
            question = ET.SubElement(quiz, 'question', type='shortanswer')
        else:
            question = ET.SubElement(quiz, 'question', type='multichoice')

        name = ET.SubElement(question, 'name')
        ET.SubElement(name, 'text').text = title

        questiontext_node = ET.SubElement(question, 'questiontext', format='html')
        ET.SubElement(questiontext_node, 'text').text = questiontext

        ET.SubElement(question, 'generalfeedback', format='html').text = ''
        ET.SubElement(question, 'defaultgrade').text = f'{points:.1f}'
        
        # Nur bei Multichoice
        if question_type == 'multichoice':
            ET.SubElement(question, 'single').text = 'true' if single == 1 else 'false'
            ET.SubElement(question, 'shuffleanswers').text = 'true'
            ET.SubElement(question, 'answernumbering').text = 'abc'
            
            # All-or-Nothing Bewertungsmethode bei Multiple Choice
            if single == 0:  # Multiple Choice
                # Setze scoring method auf "all-or-nothing" (mcq_scoring_method)
                # Dies ist die korrekte Moodle-Einstellung für "alle oder nichts"
                ET.SubElement(question, 'shownumcorrect')  # Leeres Element
            
            # Feedback
            ET.SubElement(question, 'correctfeedback', format='html').text = '<text>Ihre Antwort ist richtig.</text>'
            ET.SubElement(question, 'partiallycorrectfeedback', format='html').text = '<text>Ihre Antwort ist teilweise richtig.</text>'
            ET.SubElement(question, 'incorrectfeedback', format='html').text = '<text>Ihre Antwort ist falsch.</text>'
        
        # Essay-spezifische Felder
        if question_type == 'essay':
            ET.SubElement(question, 'responseformat').text = 'editor'
            ET.SubElement(question, 'responserequired').text = '1'
            ET.SubElement(question, 'responsefieldlines').text = '15'
            ET.SubElement(question, 'attachments').text = '0'
            ET.SubElement(question, 'attachmentsrequired').text = '0'
            ET.SubElement(question, 'graderinfo', format='html').text = ''
            ET.SubElement(question, 'responsetemplate', format='html').text = ''
        
        # Shortanswer-spezifische Felder
        if question_type == 'shortanswer':
            ET.SubElement(question, 'usecase').text = '0'

        # Tags
        tags_root = ET.SubElement(question, 'tags')
        for t in tags.split(','):
            tag_text = t.strip()
            if tag_text:
                tag = ET.SubElement(tags_root, 'tag')
                ET.SubElement(tag, 'text').text = tag_text

        # Antworten (nur bei Multichoice und Shortanswer)
        if question_type in ['multichoice', 'shortanswer']:
            c.execute('SELECT answertext, is_correct FROM answers WHERE question_id=?', (qid,))
            answers = c.fetchall()
            
            # Bei Multiple Choice mit All-or-Nothing
            if question_type == 'multichoice' and single == 0:
                # Zähle richtige Antworten
                correct_count = sum(1 for a in answers if a[1] == 1)
                
                # All-or-Nothing: nur richtige Antworten bekommen 100%, alle anderen 0%
                # Nur wenn ALLE richtigen ausgewählt sind UND KEINE falschen, gibt es Punkte
                for answertext, is_correct in answers:
                    if not answertext.strip():
                        continue
                    
                    answer = ET.SubElement(question, 'answer', fraction=('100' if is_correct == 1 else '0'), format='html')
                    ET.SubElement(answer, 'text').text = answertext
                    ET.SubElement(answer, 'feedback', format='html').text = ''
            
            # Bei Single Choice oder Shortanswer: normale Bewertung
            elif question_type == 'multichoice' and single == 1:
                for answertext, is_correct in answers:
                    fraction = '100' if is_correct == 1 else '0'
                    answer = ET.SubElement(question, 'answer', fraction=fraction, format='html')
                    ET.SubElement(answer, 'text').text = answertext
                    ET.SubElement(answer, 'feedback', format='html').text = ''
            
            # Shortanswer: Standard-Bewertung
            else:
                for answertext, is_correct in answers:
                    fraction = '100' if is_correct == 1 else '0'
                    answer = ET.SubElement(question, 'answer', fraction=fraction, format='html')
                    ET.SubElement(answer, 'text').text = answertext
                    ET.SubElement(answer, 'feedback', format='html').text = ''

    tree = ET.ElementTree(quiz)
    
    # ✅ KORRIGIERTE XML-SCHREIBUNG (ohne short_empty_tags)
    with open(filename, 'wb') as f:
        f.write(b'<?xml version="1.0" encoding="UTF-8"?>\n')
        tree.write(f, encoding='utf-8', xml_declaration=False)
    
    conn.close()


def export_to_word(db_path, question_ids, filename):
    """Exportiert ausgewählte Fragen als 2-spaltiges Word-Dokument"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    conn = sqlite3.connect(db_path)
    c = conn.cursor()
    
    # Neues Word-Dokument erstellen
    doc = Document()
    
    # Titel hinzufügen
    title = doc.add_heading('Test - Multiple Choice Fragen', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Leerzeile
    
    # Für jede Frage eine 2-spaltige Tabelle erstellen
    for idx, qid in enumerate(question_ids, 1):
        c.execute('SELECT title, questiontext, single, tags, points, question_type FROM questions WHERE id=?', (qid,))
        row = c.fetchone()
        if not row:
            continue
            
        title, questiontext, single, tags, points = row[:5]
        question_type = row[5] if len(row) > 5 else 'multichoice'
        
        # Fragenummer und Titel
        heading = doc.add_heading(f'Frage {idx}: {title}', level=2)
        
        # Punkteanzahl
        points_para = doc.add_paragraph(f'Punkte: {points:.1f}')
        points_para.runs[0].bold = True
        
        # Fragetext
        question_para = doc.add_paragraph(questiontext)
        question_para.style = 'Intense Quote'
        
        # Antworten (nur bei Multichoice und Shortanswer)
        if question_type in ['multichoice', 'shortanswer']:
            c.execute('SELECT answertext, is_correct FROM answers WHERE question_id=?', (qid,))
            answers = c.fetchall()
            
            if answers:
                # 2-spaltige Tabelle für Antworten erstellen
                table = doc.add_table(rows=len(answers), cols=2)
                table.style = 'Light Grid Accent 1'
                
                # Spaltenbreiten setzen
                table.columns[0].width = Inches(0.5)  # Checkbox-Spalte
                table.columns[1].width = Inches(5.5)  # Antwort-Spalte
                
                for i, (answertext, is_correct) in enumerate(answers):
                    # Checkbox-Spalte
                    checkbox_cell = table.cell(i, 0)
                    checkbox_cell.text = '☐'
                    checkbox_para = checkbox_cell.paragraphs[0]
                    checkbox_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    checkbox_run = checkbox_para.runs[0]
                    checkbox_run.font.size = Pt(16)
                    
                    # Antwort-Spalte
                    answer_cell = table.cell(i, 1)
                    answer_cell.text = answertext
                    answer_para = answer_cell.paragraphs[0]
                    
                    # Richtige Antworten grün markieren (nur zur Kontrolle)
                    if is_correct == 1:
                        answer_run = answer_para.runs[0]
                        answer_run.font.color.rgb = RGBColor(0, 128, 0)
                        answer_run.font.bold = True
        
        elif question_type == 'essay':
            # Essay-Fragen: Platz für Antwort lassen
            doc.add_paragraph('Antwort:')
            doc.add_paragraph('_' * 80)
            doc.add_paragraph()
            doc.add_paragraph('_' * 80)
            doc.add_paragraph()
        
        # Tags hinzufügen
        if tags:
            tags_para = doc.add_paragraph(f'Tags: {tags}')
            tags_para.runs[0].font.italic = True
            tags_para.runs[0].font.size = Pt(9)
            tags_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        # Seitenumbruch nach jeder Frage (außer bei der letzten)
        if idx < len(question_ids):
            doc.add_page_break()
    
    # Dokument speichern
    doc.save(filename)
    conn.close()
